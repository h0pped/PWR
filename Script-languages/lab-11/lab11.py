import os
import sys
import urllib
import urllib.request
import re
import math
import matplotlib.pyplot as plt
from PIL import Image
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.run import *


def extract_info(url):
    titleRegex = r'Title: (.+)\r'
    authorRegex = r'Author: (.+)\r'
    chapter1 = ""
    file = urllib.request.urlopen(url)
    decoded = file.read().decode('UTF-8')
    titleRes = re.search(titleRegex, decoded)
    authorRes = re.search(authorRegex, decoded)
    print(f"Title: {titleRes.group(1)}")
    print(f"Author: {authorRes.group(1)}")
    chapter1 = ""
    start = False
    for line in decoded.splitlines():
        if start == False:
            if line == "CHAPTER I":
                start = True
        else:
            if line != "CHAPTER II":
                chapter1 += line
                chapter1 += "\n"
            else:
                break
    res = dict()
    res["title"] = titleRes.group(1)
    res["author"] = authorRes.group(1)
    res.update(chapter_1(chapter1))
    print(res)
    return res


def chapter_1(chapter):
    words = dict()
    info = dict()
    x = 1
    counter = 0
    for line in chapter.strip().splitlines():
        if len(line) != 0:
            if x in words:
                words[x] += len(line.split())
            else:
                words[x] = len(line.split())
        else:
            x += 1
    info['first_chapter_max_words_paragraph'] = max(words, key=words.get)
    info['first_chapter_average_words'] = round(
        sum(words.values()) / len(words))
    for x in words.keys():
        words[x] = int(math.floor(words[x] / 10.0)) * 10
    # Sort by number of words
    words = dict(sorted(words.items(), key=lambda item: item[1]))
    print(words)
    same_count = []
    vals = words.values()
    for key, value in words.items():
        if list(vals).count(value) > 1:
            same_count.append(key)

    plot(words.keys(), words.values())
    info["paragraphs_count"] = len(words.keys())
    info["same_count_of_words_paragraphs_count"] = len(same_count)
    info['first_chapter_words_count'] = len(
        chapter.strip().replace("\n", "").split())
    return info


def plot(x, y):
    plt.ylabel('Count of words')
    plt.xlabel('Paragraph')
    plt.title('Chart')
    plt.figure(figsize=(10, 5))
    plt.bar(x, y)
    plt.savefig('plot.png')


def picture():
    urllib.request.urlretrieve("https://d1kvkzjpuym02z.cloudfront.net/5ebafaf4f8c55107773d91a9.jpeg?Expires=2062430380&Signature=MCgknWO053cnGWMUszWDbEhUfmd8p8OuZsHly~YYyO5ev0~OLDuThrsm9zbfArAj3eW0Qkd0CG1Vj7mut9LefTE6vVAbeQfK59ffQ2zYpE972IEV2ldUtUwH5VtQUYEIJCjNrl2oLR23a8igTFByvlLvzR3Z2q6ENeiYr5rYKCw_&Key-Pair-Id=APKAJXYWFXCDTRLR3EFA", "image1.jpeg")
    image = Image.open('image1.jpeg')
    new_image = image.resize((800, 550))
    box = (0, 0, 700, 450)
    cropped_image = new_image.crop(box)
    cropped_image.save("image1.jpeg")

    book_img = Image.open("image2.jpeg").convert("RGBA")
    book_img = book_img.resize((120, 200)).rotate(20, expand=True)
    cropped_image_copy = cropped_image.copy()
    position = ((cropped_image.width-book_img.width-10),
                (cropped_image.height-book_img.height-10))
    cropped_image_copy.paste(book_img, position, book_img)
    cropped_image_copy.save("result.jpeg")


def generate_docx(info):
    doc = docx.Document()
    p_title = doc.add_paragraph(f"Python lab #11", "Title")
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Title: {info['title']}", "Title")
    doc.add_picture('result.jpeg', width=Inches(5.0))
    doc.add_paragraph(f"Author: {info['author']}", "Title")
    doc.add_paragraph(f"Author of the report: Nykonchuk Illia", "Title")

    doc.paragraphs[3].runs[0].bold = True
    doc.paragraphs[4].runs[0].italic = True
    doc.paragraphs[4].runs[0].add_break(
        docx.text.run.WD_BREAK.PAGE)
    doc.add_paragraph(f"Chart Page", "Title")
    doc.add_paragraph(
        f"Plot about distribution of words among paragraphs in Chapter I:", "Subtitle")
    doc.paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture('plot.png', width=Inches(6.0))
    doc.add_paragraph(f"Chapter I Analysis", "Title")
    doc.paragraphs[7].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Number of paragraphs: {info['paragraphs_count']}", "Subtitle")
    doc.add_paragraph(
        f"Count of words: {info['first_chapter_words_count']}", "Subtitle")
    doc.add_paragraph(
        f"Biggest amount of words in paragraph: {info['first_chapter_max_words_paragraph']}", "Subtitle")
    doc.add_paragraph(
        f"Average amount of words per paragraph: {info['first_chapter_average_words']}", "Subtitle")
    doc.add_paragraph(
        f"Amount of paragraphs with the same amount of words: {info['same_count_of_words_paragraphs_count']}", "Subtitle")

    doc.save('myfirstdoc.docx')


def run():
    info = extract_info("https://www.gutenberg.org/files/1074/1074-0.txt")
    picture()
    generate_docx(info)


if __name__ == "__main__":
    run()

# https://www.gutenberg.org/files/1074/1074-0.txt
